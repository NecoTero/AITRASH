from __future__ import annotations

import json
import re
import sys
import zipfile
from collections import OrderedDict
from datetime import date
from pathlib import Path
from xml.etree import ElementTree as ET


WORKSPACE = Path(__file__).resolve().parents[2]
DEPS = WORKSPACE / ".tmp" / "eir_pydeps"
sys.path.insert(0, str(DEPS))

from docx import Document  # noqa: E402
from docx.enum.section import WD_ORIENT, WD_SECTION  # noqa: E402
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Inches, Pt, RGBColor  # noqa: E402


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "Обновление EIR.xlsx"
OUTPUT = ROOT / "EIR - первая сборка из базы.docx"
REPORT = WORKSPACE / ".tmp" / "eir_word_build_report.json"

NS_MAIN = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
NS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_PKG_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
NS = {"x": NS_MAIN, "r": NS_REL, "pr": NS_PKG_REL}

NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
PALE_BLUE = "EEF4FA"
GRAY = "6B7280"
LIGHT_GRAY = "E7E9ED"
WHITE = "FFFFFF"
BLACK = "1F2933"


def col_letters(ref: str) -> str:
    match = re.match(r"([A-Z]+)", ref)
    return match.group(1) if match else ""


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, inches: float) -> None:
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(round(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_and_no_split(table) -> None:
    if table.rows:
        set_repeat_table_header(table.rows[0])
    for row in table.rows:
        keep_row_together(row)


def set_run_font(run, name="Arial", size=Pt(10.5), color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_field(paragraph, instruction: str, fallback: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = fallback
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text, end))


class WorkbookReader:
    def __init__(self, path: Path):
        self.path = path
        self.zip = zipfile.ZipFile(path)
        self.shared = self._read_shared_strings()
        self.sheet_parts = self._read_sheet_parts()

    def close(self):
        self.zip.close()

    def _xml(self, name: str):
        return ET.fromstring(self.zip.read(name))

    def _read_shared_strings(self) -> list[str]:
        if "xl/sharedStrings.xml" not in self.zip.namelist():
            return []
        root = self._xml("xl/sharedStrings.xml")
        values = []
        for si in root.findall("x:si", NS):
            values.append("".join(t.text or "" for t in si.findall(".//x:t", NS)))
        return values

    def _read_sheet_parts(self) -> dict[str, str]:
        workbook = self._xml("xl/workbook.xml")
        rels = self._xml("xl/_rels/workbook.xml.rels")
        rel_map = {r.attrib["Id"]: r.attrib["Target"] for r in rels.findall("pr:Relationship", NS)}
        out = {}
        for sheet in workbook.findall("x:sheets/x:sheet", NS):
            rid = sheet.attrib[f"{{{NS_REL}}}id"]
            target = rel_map[rid].replace("\\", "/").lstrip("/")
            if not target.startswith("xl/"):
                target = "xl/" + target
            out[sheet.attrib["name"]] = target
        return out

    def _cell_value(self, cell) -> str:
        value = cell.find("x:v", NS)
        cell_type = cell.attrib.get("t", "")
        if cell_type == "s" and value is not None and value.text is not None:
            return self.shared[int(value.text)]
        if cell_type == "inlineStr":
            return "".join(t.text or "" for t in cell.findall(".//x:t", NS))
        if value is None or value.text is None:
            return ""
        return value.text

    def sheet_matrix(self, name: str) -> list[dict[str, str]]:
        root = self._xml(self.sheet_parts[name])
        result = []
        for row in root.findall("x:sheetData/x:row", NS):
            values = OrderedDict()
            values["__row__"] = row.attrib.get("r", "")
            for cell in row.findall("x:c", NS):
                values[col_letters(cell.attrib["r"])] = self._cell_value(cell)
            result.append(values)
        return result

    def records(self, name: str, header_row: int = 3) -> list[dict[str, str]]:
        matrix = self.sheet_matrix(name)
        headers = next(row for row in matrix if row.get("__row__") == str(header_row))
        col_to_header = {col: value.strip() for col, value in headers.items() if col != "__row__" and value.strip()}
        records = []
        for raw in matrix:
            if int(raw.get("__row__", "0") or 0) <= header_row:
                continue
            record = {header: raw.get(col, "").strip() for col, header in col_to_header.items()}
            if any(record.values()):
                records.append(record)
        return records


def set_style_font(style, size: float, bold=False, color=BLACK, name="Arial") -> None:
    style.font.name = name
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title, 26, bold=True, color=NAVY)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    for name, size, color in (("Heading 1", 16, NAVY), ("Heading 2", 13, BLUE), ("Heading 3", 11, BLUE)):
        style = styles[name]
        set_style_font(style, size, bold=True, color=color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    styles["Heading 1"].paragraph_format.page_break_before = True
    styles["Heading 1"].paragraph_format.keep_with_next = True

    subtitle = styles["Subtitle"]
    set_style_font(subtitle, 12, color=GRAY)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)

    if "Requirement" not in styles:
        styles.add_style("Requirement", 1)
    requirement = styles["Requirement"]
    set_style_font(requirement, 10.5)
    requirement.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    requirement.paragraph_format.left_indent = Cm(0.75)
    requirement.paragraph_format.first_line_indent = Cm(-0.75)
    requirement.paragraph_format.space_after = Pt(6)
    requirement.paragraph_format.line_spacing = 1.12
    requirement.paragraph_format.widow_control = True

    if "EIR Bullet" not in styles:
        styles.add_style("EIR Bullet", 1)
    bullet = styles["EIR Bullet"]
    set_style_font(bullet, 10.25)
    bullet.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bullet.paragraph_format.left_indent = Cm(1.25)
    bullet.paragraph_format.first_line_indent = Cm(-0.55)
    bullet.paragraph_format.space_after = Pt(3)
    bullet.paragraph_format.line_spacing = 1.08

    if "Table Caption" not in styles:
        styles.add_style("Table Caption", 1)
    caption = styles["Table Caption"]
    set_style_font(caption, 9.5, bold=True, color=NAVY)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.keep_with_next = True


def configure_section(section, landscape=False) -> None:
    section.different_first_page_header_footer = False
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width, section.page_height = Cm(29.7), Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.7)
        section.bottom_margin = Cm(1.7)
    else:
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(1.8)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)


def add_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = "EIR  |  РАБОЧАЯ РЕДАКЦИЯ"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        set_run_font(run, size=Pt(8), color=GRAY, bold=True)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LIGHT_GRAY)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Страница ")
    set_run_font(run, size=Pt(8), color=GRAY)
    add_field(p, "PAGE", "1")
    run = p.add_run(" из ")
    set_run_font(run, size=Pt(8), color=GRAY)
    add_field(p, "NUMPAGES", "1")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Cm(3.3)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ИНФОРМАЦИОННЫЕ\nТРЕБОВАНИЯ ЗАКАЗЧИКА")
    set_run_font(r, size=Pt(24), color=NAVY, bold=True)
    r.add_break()
    r2 = p.add_run("Employer's Information Requirements (EIR)")
    set_run_font(r2, size=Pt(14), color=BLUE, bold=True)

    p = doc.add_paragraph(style="Subtitle")
    p.add_run("Рабочая редакция\nПервая автоматизированная сборка из базы требований")

    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line.paragraph_format.space_before = Cm(1.2)
    r = line.add_run("────────────────────────")
    set_run_font(r, size=Pt(11), color=LIGHT_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Cm(3.4)
    r = meta.add_run("Источник данных: Обновление EIR.xlsx\n")
    set_run_font(r, size=Pt(10), color=GRAY)
    r = meta.add_run(f"Дата сборки: {date.today().strftime('%d.%m.%Y')}")
    set_run_font(r, size=Pt(10), color=GRAY)

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_toc(doc: Document) -> None:
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(12)
    r = heading.add_run("СОДЕРЖАНИЕ")
    set_run_font(r, size=Pt(17), color=NAVY, bold=True)
    field_p = doc.add_paragraph()
    add_field(field_p, 'TOC \\o "1-1" \\h \\z \\u', "Оглавление обновляется при открытии документа")
    field_p.paragraph_format.space_after = Pt(6)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = note.add_run("Документ сформирован из публикуемых полей базы EIR. Служебные заметки, внутренние ID и поля проверки в основной текст не включены.")
    set_run_font(r, size=Pt(8.5), color=GRAY, italic=True)


def add_requirement(doc: Document, record: dict[str, str]) -> None:
    p = doc.add_paragraph(style="Requirement")
    p.paragraph_format.keep_together = True
    number = record.get("Номер_для_публикации", "")
    text = record.get("Формулировка_требования", "")
    text_lines = [line.strip() for line in text.splitlines()]
    first_line = text_lines[0] if text_lines else ""
    if number:
        r = p.add_run(number + "\u00a0\u00a0")
        set_run_font(r, size=Pt(10.5), color=NAVY, bold=True)
    r = p.add_run(first_line)
    set_run_font(r, size=Pt(10.5), color=BLACK)
    for line in (line for line in text_lines[1:] if line):
        child = doc.add_paragraph(style="EIR Bullet")
        child.paragraph_format.keep_together = True
        r = child.add_run("–\u00a0\u00a0")
        set_run_font(r, size=Pt(10.25), color=BLUE, bold=True)
        r = child.add_run(line)
        set_run_font(r, size=Pt(10.25), color=BLACK)


def add_bullet(doc: Document, record: dict[str, str]) -> None:
    p = doc.add_paragraph(style="EIR Bullet")
    p.paragraph_format.keep_together = True
    depth = max(0, int(float(record.get("Уровень", "3") or 3)) - 3)
    p.paragraph_format.left_indent = Cm(1.25 + depth * 0.65)
    p.paragraph_format.first_line_indent = Cm(-0.55)
    marker = record.get("Маркер_перечня", "")
    bullet_char = "–" if marker else "•"
    r = p.add_run(bullet_char + "\u00a0\u00a0")
    set_run_font(r, size=Pt(10.25), color=BLUE, bold=True)
    title = record.get("Заголовок_элемента", "")
    text = record.get("Формулировка_требования", "")
    if title:
        r = p.add_run(title)
        set_run_font(r, size=Pt(10.25), color=BLACK, bold=True)
        if text:
            separator = ". " if not title.rstrip().endswith((".", ":", ";")) else " "
            r = p.add_run(separator)
            set_run_font(r, size=Pt(10.25), color=BLACK)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=Pt(10.25), color=BLACK)


def add_subheading(doc: Document, record: dict[str, str]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    text = record.get("Формулировка_требования", "") or record.get("Заголовок_элемента", "")
    r = p.add_run(text)
    set_run_font(r, size=Pt(10.5), color=BLUE, bold=True, italic=True)


def table_caption(doc: Document, number: int, table_id: str, title: str) -> None:
    p = doc.add_paragraph(style="Table Caption")
    r = p.add_run(f"Таблица {number} ({table_id}) — {title}")
    set_run_font(r, size=Pt(9.5), color=NAVY, bold=True)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size=8.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, top=110, start=90, bottom=110, end=90)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=Pt(font_size), color=WHITE, bold=True)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            cell.text = value
            set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if row_idx % 2 == 1:
                set_cell_shading(cell, PALE_BLUE)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 or value.lower() in {"x", "да", "нет"} else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_run_font(run, size=Pt(font_size), color=BLACK)
    set_repeat_and_no_split(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_public_table(doc: Document, table_id: str, table_records: dict[str, list[dict[str, str]]], counter: int) -> tuple[int, bool]:
    if table_id == "TBL-0001":
        records = [r for r in table_records[table_id] if r.get("Статус") != "Отменено"]
        records.sort(key=lambda r: float(r.get("Порядок", "0") or 0))
        table_caption(doc, counter, table_id, "Сценарии использования BIM-моделей")
        add_table(
            doc,
            ["Код", "Название", "Назначение", "Обязательность"],
            [[r.get("Код", ""), r.get("Название", ""), r.get("Назначение", ""), r.get("Обязательность", "")] for r in records],
            [0.75, 1.55, 3.65, 1.0],
            font_size=8.5,
        )
        return counter + 1, False
    if table_id == "TBL-0002":
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section, landscape=True)
        add_header_footer(section)
        records = [r for r in table_records[table_id] if r.get("Статус") != "Отменено"]
        records.sort(key=lambda r: float(r.get("Порядок", "0") or 0))
        table_caption(doc, counter, table_id, "Инструменты и параметры расчёта ТЭП")
        headers = ["Наименование ТЭП", "Инструмент", "Точный\nгабарит", "Точное\nположение", "OMDV_\nКорпус", "OMDV_\nСекция", "OMDV_\nНомер этажа", "OMDV_\nНазначение", "Обязательность"]
        fields = ["Наименование ТЭП", "Инструмент выполнения", "Точный габарит", "Точное положение", "OMDV_Корпус", "OMDV_Секция", "OMDV_Номер этажа", "OMDV_Назначение", "Обязательность"]
        add_table(doc, headers, [[r.get(f, "") for f in fields] for r in records], [1.75, 1.15, 0.78, 0.78, 0.95, 0.95, 1.1, 1.05, 0.95], font_size=7.5)
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        configure_section(section, landscape=False)
        add_header_footer(section)
        return counter + 1, True
    return counter, False


def build() -> dict:
    reader = WorkbookReader(SOURCE)
    try:
        records = reader.records("EIR БД")
        records = [r for r in records if r.get("Статус") != "Отменено"]
        records.sort(key=lambda r: float(r.get("Порядок_публикации", "0") or 0))
        table_records = {
            "TBL-0001": reader.records("TBL-0001"),
            "TBL-0002": reader.records("TBL-0002"),
        }
    finally:
        reader.close()

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0], landscape=False)
    add_header_footer(doc.sections[0])
    doc.sections[0].different_first_page_header_footer = True
    doc.sections[0].first_page_header.paragraphs[0].text = ""
    doc.sections[0].first_page_footer.paragraphs[0].text = ""
    add_cover(doc)
    add_toc(doc)

    section_titles = OrderedDict()
    for record in records:
        sid = record.get("Раздел_ID", "")
        title = record.get("Раздел_EIR", "")
        if sid and title and sid not in section_titles:
            section_titles[sid] = title

    current_section = None
    table_counter = 1
    added = 0
    formulations = []
    for record in records:
        section_id = record.get("Раздел_ID", "")
        if section_id != current_section:
            current_section = section_id
            title = section_titles.get(section_id, f"Раздел {section_id}")
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(f"{section_id}. {title}")
            set_run_font(r, size=Pt(16), color=NAVY, bold=True)

        record_type = record.get("Тип_записи", "")
        text = record.get("Формулировка_требования", "")
        if text:
            formulations.append(text)
        if record_type in {"Требование", "Ссылка_на_таблицу"}:
            add_requirement(doc, record)
            added += 1
        elif record_type == "Элемент_перечня":
            add_bullet(doc, record)
            added += 1
        elif record_type == "Заголовок":
            add_subheading(doc, record)
            added += 1
        elif text:
            add_requirement(doc, record)
            added += 1

        table_id = record.get("Таблица_ID", "")
        if table_id:
            table_counter, _ = add_public_table(doc, table_id, table_records, table_counter)

    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")

    props = doc.core_properties
    props.title = "Информационные требования Заказчика (EIR)"
    props.subject = "Первая автоматизированная сборка из базы требований"
    props.author = "ОМ Девелопмент"
    props.keywords = "EIR, BIM, информационные требования"

    doc.save(OUTPUT)

    reopened = Document(OUTPUT)
    visible = "\n".join(p.text for p in reopened.paragraphs)
    visible += "\n" + "\n".join(cell.text for table in reopened.tables for row in table.rows for cell in row.cells)
    normalize = lambda value: re.sub(r"\s+", " ", value).strip()
    visible_normalized = normalize(visible)
    def formulation_present(text: str) -> bool:
        parts = [normalize(line) for line in text.splitlines() if normalize(line)]
        return all(part in visible_normalized for part in parts)

    missing = [text for text in formulations if not formulation_present(text)]
    report = {
        "source": str(SOURCE),
        "output": str(OUTPUT),
        "records_source": len(records),
        "records_added": added,
        "formulations_expected": len(formulations),
        "formulations_missing": len(missing),
        "missing_samples": missing[:5],
        "sections": section_titles,
        "tables_inserted": table_counter - 1,
        "document_paragraphs": len(reopened.paragraphs),
        "document_tables": len(reopened.tables),
    }
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if missing:
        raise SystemExit("Source-to-document validation failed")
    return report


if __name__ == "__main__":
    build()
